import streamlit as st
import pandas as pd
from datetime import timedelta
import openpyxl
from openpyxl.styles import Alignment, Font
import io
import re
import smtplib
import urllib.parse as _ul
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import docx
import pdfplumber

try:
    from menu import show_sidebar
except ImportError:
    def show_sidebar():
        pass

def send_csv_email(file_bytes, file_name, year_str):
    try:
        sender, pwd = st.secrets["email"]["user"], st.secrets["email"]["password"]
        msg = MIMEMultipart()
        msg["From"], msg["To"] = sender, sender
        # 信件標題直接使用動態生成的完整檔名
        msg["Subject"] = f"{file_name}"
        body_text = "您好，\n\n系統已自動依據輪休預排與上傳範本產出防制危險駕車專案督勤表，附件為對應的 Excel 統計檔案。\n\n本信件由交通執法自動化分析引擎發送。"
        msg.attach(MIMEText(body_text, "plain", "utf-8"))
        part = MIMEBase("application", "vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        part.set_payload(file_bytes.getvalue())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename*=UTF-8''{_ul.quote(file_name)}")
        msg.attach(part)
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender, pwd)
            server.sendmail(sender, sender, msg.as_string())
        return True, None
    except Exception as e:
        return False, str(e)

def clean_text(text):
    if not text:
        return ""
    return re.sub(r'\s+', '', str(text))

def extract_vacations_from_files(uploaded_vacations_list):
    vacations = {}
    if not uploaded_vacations_list:
        return vacations
        
    for uploaded_file in uploaded_vacations_list:
        try:
            month = "1"
            month_match = re.search(r'(\d+)月', uploaded_file.name)
            if month_match:
                month = str(int(month_match.group(1)))
                
            ext = uploaded_file.name.lower().split('.')[-1]

            if ext == "pdf":
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text() or ""
                        m_match = re.search(r'年\s*(\d+)\s*月', text)
                        if m_match:
                            month = str(int(m_match.group(1)))
                            
                        # 強制使用垂直水平線條辨識策略
                        tables = page.extract_tables(table_settings={"vertical_strategy": "text", "horizontal_strategy": "text"})
                        if not tables:
                            tables = page.extract_tables() # fallback
                            
                        for table in tables:
                            col_names = {}
                            start_row_idx = 0
                            
                            for i, row in enumerate(table):
                                cells_text = [clean_text(c) if c else "" for c in row]
                                if any("分局長" in c for c in cells_text):
                                    for col_idx, txt in enumerate(cells_text):
                                        if txt and "日期" not in txt and "星期" not in txt and "輪休" not in txt:
                                            col_names[col_idx] = txt
                                    start_row_idx = i + 1
                                    break
                            
                            if not col_names: continue
                            
                            for i in range(start_row_idx, len(table)):
                                cells_text = [clean_text(c) if c else "" for c in table[i]]
                                if not cells_text or not cells_text[0].isdigit():
                                    continue
                                day = int(cells_text[0])
                                date_val = f"{month}/{day}"
                                
                                for col_idx, officer_name in col_names.items():
                                    if col_idx < len(cells_text):
                                        mark = cells_text[col_idx]
                                        if mark and mark.strip() != "":
                                            vacations.setdefault(officer_name, []).append(date_val)
                                            
            elif ext in ["docx", "doc"]:
                doc = docx.Document(uploaded_file)
                if doc.paragraphs:
                    m_match = re.search(r'年\s*(\d+)\s*月', doc.paragraphs[0].text)
                    if m_match:
                        month = str(int(m_match.group(1)))
                        
                for table in doc.tables:
                    col_names = {}
                    start_row_idx = 0
                    for i, row in enumerate(table.rows):
                        cells_text = [clean_text(cell.text) for cell in row.cells]
                        if any("分局長" in c for c in cells_text):
                            for col_idx, txt in enumerate(cells_text):
                                if txt and "日期" not in txt and "星期" not in txt and "輪休" not in txt:
                                    col_names[col_idx] = txt
                            start_row_idx = i + 1
                            break
                    if not col_names: continue
                    for i in range(start_row_idx, len(table.rows)):
                        cells_text = [clean_text(cell.text) for cell in table.rows[i].cells]
                        if not cells_text or not cells_text[0].isdigit(): continue
                        day = int(cells_text[0])
                        date_val = f"{month}/{day}"
                        for col_idx, officer_name in col_names.items():
                            if col_idx < len(cells_text):
                                mark = cells_text[col_idx]
                                if mark and mark.strip() != "":
                                    vacations.setdefault(officer_name, []).append(date_val)
        except Exception as e:
            st.warning(f"休假表「{uploaded_file.name}」解析失敗：{e}")
            
    for k in vacations: vacations[k] = list(set(vacations[k]))
    return vacations

def process_excel(uploaded_file, df_personnel, full_date_list, hours_per_shift, vacation_dict):
    wb = openpyxl.load_workbook(uploaded_file)
    ws = wb['警力統計'] if '警力統計' in wb.sheetnames else wb.active
    title_cell = ws['A1']
    if title_cell.value and '○○分局' in title_cell.value:
        title_cell.value = title_cell.value.replace('○○分局', '龍潭分局')

    # 先取得排休表上有哪些副分局長，確保按照欄位順序排序
    deputy_list = [k for k in vacation_dict.keys() if "副分局長" in k]

    current_row = 3
    deputy_counter = 0 # 紀錄目前處理到第幾個副分局長

    for _, row in df_personnel.iterrows():
        title = clean_text(row.get("職稱", ""))
        name = clean_text(row.get("姓名", ""))
        raw_title = str(row.get("職稱", "")).strip()
        raw_name = str(row.get("姓名", "")).strip()
        
        if raw_title or raw_name:
            ws.cell(row=current_row, column=1, value=raw_title)
            ws.cell(row=current_row, column=2, value=raw_name)
            
            user_vacations = []
            
            # 【智慧與防呆對位邏輯】
            if title == "分局長":
                for v_name, dates in vacation_dict.items():
                    if "分局長" in v_name and "副" not in v_name:
                        user_vacations.extend(dates)
            
            elif "副分局長" in title:
                matched = False
                # 先嘗試用姓名對位 (例如 姓名有"何")
                if name:
                    for v_name, dates in vacation_dict.items():
                        if name[0] in v_name:
                            user_vacations.extend(dates)
                            matched = True
                # 如果沒填名字，或是名字沒對到，自動依照順序給予休假資料
                if not matched and deputy_counter < len(deputy_list):
                    v_name = deputy_list[deputy_counter]
                    user_vacations.extend(vacation_dict[v_name])
                deputy_counter += 1
            
            # 過濾休假日期
            person_dates = []
            for date_str in full_date_list:
                raw_date = date_str.split('(')[0]
                if raw_date not in user_vacations:
                    person_dates.append(date_str)
                    
            cell_c = ws.cell(row=current_row, column=3)
            cell_c.value = "、".join(person_dates)
            cell_c.alignment = Alignment(wrapText=True, vertical='center', horizontal='left')
            
            cell_d = ws.cell(row=current_row, column=4)
            cell_d.value = len(person_dates) * hours_per_shift
            cell_d.alignment = Alignment(vertical='center', horizontal='center')
            cell_d.font = Font(name='微軟正黑體', size=12)
            
            current_row += 1

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def main():
    show_sidebar()
    st.title("🗓️ 上半年連假專案督勤表生成")
    st.info("請上傳空白範本與輪休表，系統將自動解析 PDF 並排除休假人員。")

    col1, col2, col3 = st.columns(3)
    with col1: uploaded_pdf = st.file_uploader("1. 辦公日曆表 (PDF)", type=['pdf'], key="p25_pdf")
    with col2: uploaded_excel = st.file_uploader("2. Excel 空白範本", type=['xlsx'], key="p25_excel")
    with col3: uploaded_vacations = st.file_uploader("3. 輪休預排表 (多檔 PDF/Word)", type=['pdf', 'docx', 'doc'], accept_multiple_files=True, key="p25_vac")

    default_personnel = pd.DataFrame([
        {"職稱": "分局長", "姓名": ""},
        {"職稱": "副分局長(一)", "姓名": ""},
        {"職稱": "副分局長(二)", "姓名": ""},
        {"職稱": "交通組組長", "姓名": ""}
    ])

    st.subheader("👥 督導人員名單設定 (防呆已啟動：不填姓名亦可自動對位)")
    edited_personnel = st.data_editor(default_personnel, num_rows="dynamic", use_container_width=True, key="p25_editor")

    year_str, ce_year = "115年", 2026
    if uploaded_excel:
        wb_temp = openpyxl.load_workbook(uploaded_excel, read_only=True)
        ws_temp = wb_temp['警力統計'] if '警力統計' in wb_temp.sheetnames else wb_temp.active
        match = re.search(r'(\d{2,3})年', ws_temp['A1'].value or "")
        if match:
            year_str, ce_year = f"{match.group(1)}年", int(match.group(1)) + 1911

    # 假日運算
    holidays_H1 = [
        {"start": f"{ce_year}-01-01", "end": f"{ce_year}-01-01"},
        {"start": f"{ce_year}-01-03", "end": f"{ce_year}-01-04"},
        {"start": f"{ce_year}-01-10", "end": f"{ce_year}-01-11"},
        {"start": f"{ce_year}-01-17", "end": f"{ce_year}-01-18"},
        {"start": f"{ce_year}-01-24", "end": f"{ce_year}-01-25"},
        {"start": f"{ce_year}-01-31", "end": f"{ce_year}-02-01"},
        {"start": f"{ce_year}-02-07", "end": f"{ce_year}-02-08"},
        {"start": f"{ce_year}-02-14", "end": f"{ce_year}-02-22"},
        {"start": f"{ce_year}-02-27", "end": f"{ce_year}-03-01"},
        {"start": f"{ce_year}-03-07", "end": f"{ce_year}-03-08"},
        {"start": f"{ce_year}-03-14", "end": f"{ce_year}-03-15"},
        {"start": f"{ce_year}-03-21", "end": f"{ce_year}-03-22"},
        {"start": f"{ce_year}-03-28", "end": f"{ce_year}-03-29"},
        {"start": f"{ce_year}-04-03", "end": f"{ce_year}-04-06"},
        {"start": f"{ce_year}-04-11", "end": f"{ce_year}-04-12"},
        {"start": f"{ce_year}-04-18", "end": f"{ce_year}-04-19"},
        {"start": f"{ce_year}-04-25", "end": f"{ce_year}-04-26"},
        {"start": f"{ce_year}-05-01", "end": f"{ce_year}-05-03"},
        {"start": f"{ce_year}-05-09", "end": f"{ce_year}-05-10"},
        {"start": f"{ce_year}-05-16", "end": f"{ce_year}-05-17"},
        {"start": f"{ce_year}-05-23", "end": f"{ce_year}-05-24"},
        {"start": f"{ce_year}-05-30", "end": f"{ce_year}-05-31"},
        {"start": f"{ce_year}-06-06", "end": f"{ce_year}-06-07"},
        {"start": f"{ce_year}-06-13", "end": f"{ce_year}-06-14"},
        {"start": f"{ce_year}-06-19", "end": f"{ce_year}-06-21"},
        {"start": f"{ce_year}-06-27", "end": f"{ce_year}-06-28"}
    ]

    full_date_list = []
    active_months = set() # 紀錄運算涵蓋的月份
    
    for h in holidays_H1:
        date_range = pd.date_range(start=pd.to_datetime(h["start"]) - timedelta(days=1), end=pd.to_datetime(h["end"]) - timedelta(days=1))
        for d in date_range: 
            full_date_list.append(f"{d.month}/{d.day}(22-06)")
            active_months.add(d.month)
            
    st.divider()
    st.subheader("📥 匯出與寄送督勤表")
    
    if uploaded_excel:
        try:
            vacation_dict = extract_vacations_from_files(uploaded_vacations)
            excel_data = process_excel(uploaded_excel, edited_personnel, full_date_list, 8, vacation_dict)
            
            # 🔥 自動依據涵蓋的月份決定檔名中的「區間」
            if active_months:
                min_m, max_m = min(active_months), max(active_months)
                period_str = f"{min_m}至{max_m}月" if min_m != max_m else f"{min_m}月"
            else:
                period_str = "未定期間"
                
            # 正式組合標準檔名
            file_name = f"桃園市政府警察局龍潭分局{year_str}{period_str}6序列以上人員督導防制危險駕車勤務時數統計表.xlsx"

            if uploaded_vacations:
                st.success(f"✅ 已掛載 {len(uploaded_vacations)} 份輪休預排表！")
                with st.expander("🔍 點此查看系統抓取的長官休假清單（除錯用，空陣列代表沒抓到假）"):
                    st.write(vacation_dict)

            b1, b2 = st.columns(2)
            with b1:
                st.download_button("📥 下載督勤日期表 (Excel)", excel_data, file_name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            with b2:
                if st.button("📧 將此報表一鍵寄至我的信箱", use_container_width=True, key="p25_btn"):
                    with st.spinner("信件發送中，請稍候…"):
                        ok, mail_err = send_csv_email(excel_data, file_name, year_str)
                        if ok: st.success("✅ 信件發送成功！")
                        else: st.error(f"❌ 發信失敗: {mail_err}")
        except Exception as e:
            st.error(f"⚠️ 處理檔案時發生錯誤：{e}")
    else:
        st.warning("⚠️ 請先上傳 Excel 空白範本，方可進行下載與寄送。")

if __name__ == "__main__":
    main()
